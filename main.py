import os
import csv
import urllib.request
import urllib.parse
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
import smtplib
from email.message import EmailMessage

# =========================
# 1. 基本設定
# =========================
month = datetime.now().strftime("%Y_%m")
csv_file = f"news_{month}.csv"
word_file = f"report_instruction_{month}.docx"

email_user = os.getenv("EMAIL_USER")
email_pass = os.getenv("EMAIL_PASS")
email_to = os.getenv("EMAIL_TO")

# 🌍 升級雙語雙軌搜尋：針對不同需求切換語系
TARGETS = [
    {"company": "Oracle", "keyword": "Oracle cloud AI server data center", "lang": "en"},
    {"company": "Oracle", "keyword": "甲骨文 伺服器 雲端", "lang": "tw"},
    {"company": "Wiwynn", "keyword": "Wiwynn AI server", "lang": "en"},
    {"company": "Wiwynn", "keyword": "緯穎 AI 伺服器 散熱", "lang": "tw"}
]

# =========================
# 2. 自動上網搜集情報 (雙語爬蟲模組)
# =========================
print(f"🔍 開始啟動 {month} 月份情報搜集 (包含國際外媒與台灣供應鏈)...")
all_news = []

for target in TARGETS:
    company = target["company"]
    safe_keyword = urllib.parse.quote(f"{target['keyword']} when:30d")
    
    # 根據設定決定要爬美國外媒還是台灣新聞
    if target["lang"] == "en":
        url = f"https://news.google.com/rss/search?q={safe_keyword}&hl=en-US&gl=US&ceid=US:en"
    else:
        url = f"https://news.google.com/rss/search?q={safe_keyword}&hl=zh-TW&gl=TW&ceid=TW:zh-Hant"
    
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response:
            xml_data = response.read()
            
        root = ET.fromstring(xml_data)
        for item in root.findall('.//item'):
            title = item.find('title').text
            pub_date = item.find('pubDate').text
            # 統一時間格式處理
            try:
                date_str = datetime.strptime(pub_date, "%a, %d %b %Y %H:%M:%S %Z").strftime("%Y-%m-%d")
            except:
                date_str = datetime.now().strftime("%Y-%m-%d") # 若解析失敗則用今天
            all_news.append([date_str, title, company])
    except Exception as e:
        print(f"❌ 抓取 {company} ({target['lang']}) 新聞時發生錯誤: {e}")

# 去除重複的新聞標題
unique_news = []
seen_titles = set()
for news in all_news:
    if news[1] not in seen_titles:
        unique_news.append(news)
        seen_titles.add(news[1])

with open(csv_file, mode="w", newline="", encoding="utf-8") as f:
    writer = csv.writer(f)
    writer.writerow(["date", "title", "company"])
    writer.writerows(unique_news)
print(f"✅ 成功扒下 {len(unique_news)} 則跨國新聞，已存入 {csv_file}！")

# =========================
# 3. 讀取並分類新聞
# =========================
market_news, tech_news, finance_news = [], [], []

with open(csv_file, newline="", encoding="utf-8") as f:
    reader = csv.DictReader(f)
    for r in reader:
        title = r.get("title", "").lower()
        company = r.get("company", "未知公司")
        content = f"- {company}：{r.get('title', '')}"
        
        # 中英關鍵字同步擴充
        if any(k in title for k in ["cloud", "ai", "market", "demand", "growth", "雲端", "市場", "需求"]):
            market_news.append(content)
        if any(k in title for k in ["data center", "server", "rack", "infrastructure", "伺服器", "機櫃", "散熱", "液冷"]):
            tech_news.append(content)
        if any(k in title for k in ["capex", "investment", "revenue", "order", "營收", "資本", "財報", "訂單"]):
            finance_news.append(content)

# =========================
# 4. 組合 Flare 專屬「華爾街分析師」指令字串
# =========================
ai_prompt = f"""
你是一位【華爾街資深產業分析師 / 顧問公司資深策略顧問】。

請根據以下已整理的新聞資料、公司資訊與產業背景，撰寫一份【Oracle & Wiwynn 產業分析月報】。

請嚴格遵守以下規範：
- 請直接輸出一份word的完整報告
- 需參考我提供的新聞與資料，將資訊整理為重點並進一步分析，不可只是摘要重述
- 請勿提及分析步驟、推理過程或你如何整理資料
- 請使用繁體中文
- 語氣需專業、冷靜、策略導向，風格需接近 McKinsey / BCG / Deloitte / sell-side 產業研究報告
- 每一段皆需為可直接貼入 Word 的完整段落
- 避免口語化表達，內容需具備管理層簡報與投資研究可讀性
- 請以「觀點 + 分析 + 商業含義」方式撰寫，而非單純新聞整理
- 報告最後請附上新聞或資料出處

＝＝＝＝ 指定報告結構（請嚴格遵守）＝＝＝＝

【一、產業與市場趨勢】
請從 Oracle 與 Wiwynn 在 AI、雲端、資料中心與企業 IT 領域的策略定位切入，
分析市場需求變化、客戶採購趨勢、產業週期與商業模式演進，
並連結供應鏈與競爭對手的發展方向，說明其產業意涵。

【二、技術演進與基礎設施設計重點】
請聚焦關鍵技術主題
分析技術演進如何影響資料中心架構、資本支出、產品規格與未來設計方向。

【三、財務與策略觀察（含競爭格局）】
請分析 Oracle 與 Wiwynn 的資本支出、營收結構、訂單動能與策略意 
並說明 foxconn、Quanta、Wistron、Inventec、Dell、HPE、Supermicro、Aivres等競爭對手
在該市場中的產業狀況、角色變化、競爭格局，以及中長期風險與機會。

【四、結論與策略建議】
請總結本月最重要的產業訊號，
並提出未來發展的評估和分析包含資源配置、競爭應對策略與未來觀察重點
請清楚區分短期市場變化與中長期結構性趨勢，
並說明這些變化對 Oracle、ODM/OEM 與終端客戶各自的商業含義。

＝＝＝＝ 本月已整理新聞資料 (全球+在地) ＝＝＝＝

【市場趨勢相關新聞】
{chr(10).join(market_news)}

【技術更新相關新聞】
{chr(10).join(tech_news)}

【財務與策略相關新聞】
{chr(10).join(finance_news)}
"""

# =========================
# 5. 存入 Word
# =========================
doc = Document()
doc.add_heading(f"Oracle & Wiwynn 產業分析指令 - {month}", 0)
doc.add_paragraph("Flare，請全選下方內容並貼至 ChatGPT / Gemini 網頁版即可產出專業報告：")
doc.add_paragraph("--------------------------------------------------")
doc.add_paragraph(ai_prompt)
doc.save(word_file)

# =========================
# 6. 寄送郵件 (包含 Word 附件)
# =========================
if email_user and email_pass:
    print(f"📧 偵測到 Email 帳號 {email_user}，準備寄出...")
    msg = EmailMessage()
    msg["Subject"] = f"【AI 指令產出】Oracle & Wiwynn 產業分析 - {month}"
    msg["From"] = email_user
    msg["To"] = email_to
    msg.set_content(f"Flare 你好，本月的 AI 分析指令已準備好。請打開附件 Word 複製文字後貼給 AI。")
    with open(word_file, "rb") as f:
        msg.add_attachment(f.read(), maintype="application", subtype="docx", filename=word_file)
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(email_user, email_pass)
            smtp.send_message(msg)
        print("✅ 指令 Word 已成功寄出！")
    except Exception as e:
        print(f"❌ 郵件寄送失敗: {e}")
else:
    print("⚠️ 警告：未讀取到 EMAIL_USER 或 EMAIL_PASS，跳過寄信步驟。請檢查 GitHub Actions 的 .yml 設定！")
